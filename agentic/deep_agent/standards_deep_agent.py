# agentic/standards_deep_agent.py
# Standards Deep Agent - Parallel Document Analysis with Map-Reduce Pattern
#
# This agent replaces the corrective RAG system with a high-integrity "Deep Agent"
# that performs parallel analysis of full standard documents. The goal is to maximize
# the accuracy of specification extraction by allowing an LLM to "read" the relevant
# standards in the context of the user's specific requirement.
#
# ITERATIVE GENERATION: If specs count < MIN_STANDARDS_SPECS_COUNT (30), the agent
# will iterate through additional standard domains until the minimum is reached.
#
# Architecture: Map-Reduce Pattern using LangGraph
# - Planner: Decides which standards are relevant
# - Workers: Parallel LLM calls to analyze each standard document
# - Synthesizer: Merges worker outputs, detects conflicts
# - Merger: Combines with database-inferred specs
# - Iterative Loop: Re-analyze additional domains if specs < minimum

import json
import logging
import time
import os
from typing import Dict, Any, List, Optional, TypedDict, Annotated, Literal
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import operator

from langgraph.graph import StateGraph, END
from langgraph.types import Send
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from docx import Document as DocxDocument

from dotenv import load_dotenv

from ..checkpointing import compile_with_checkpointing
from ..vector_store import get_vector_store
from llm_fallback import create_llm_with_fallback

load_dotenv()
logger = logging.getLogger(__name__)


# ============================================================================
# STANDARD DOCUMENT DEFINITIONS
# ============================================================================

# Available standard document domains
STANDARD_DOMAINS = {
    "safety": {
        "name": "Safety & Protection Standards",
        "keywords": ["sil", "safety", "atex", "iecex", "hazardous", "zone", "explosion", "protection", "sis", "sif"],
        "description": "SIL ratings, ATEX zones, hazardous area classifications, safety instrumented systems"
    },
    "pressure": {
        "name": "Pressure Measurement Standards",
        "keywords": ["pressure", "transmitter", "gauge", "differential", "absolute", "psi", "bar", "pascal"],
        "description": "Pressure measurement devices, transmitters, gauges, calibration requirements"
    },
    "temperature": {
        "name": "Temperature Measurement Standards",
        "keywords": ["temperature", "thermocouple", "rtd", "thermometer", "celsius", "fahrenheit", "kelvin", "thermal"],
        "description": "Temperature sensors, thermocouples, RTDs, calibration standards"
    },
    "flow": {
        "name": "Flow Measurement Standards",
        "keywords": ["flow", "flowmeter", "coriolis", "ultrasonic", "vortex", "turbine", "mass flow", "volumetric"],
        "description": "Flow measurement devices, mass flow, volumetric flow standards"
    },
    "level": {
        "name": "Level Measurement Standards",
        "keywords": ["level", "tank", "radar", "ultrasonic", "float", "displacer", "capacitance", "hydrostatic"],
        "description": "Level measurement for tanks, vessels, silos"
    },
    "control": {
        "name": "Control Systems Standards",
        "keywords": ["control", "pid", "dcs", "plc", "scada", "controller", "loop", "feedback"],
        "description": "Process control systems, DCS, PLC, control loops"
    },
    "valves": {
        "name": "Valves & Actuators Standards",
        "keywords": ["valve", "actuator", "positioner", "control valve", "solenoid", "pneumatic", "hydraulic"],
        "description": "Control valves, actuators, positioners, valve specifications"
    },
    "calibration": {
        "name": "Calibration & Maintenance Standards",
        "keywords": ["calibration", "maintenance", "traceability", "uncertainty", "accuracy", "tolerance"],
        "description": "Calibration procedures, maintenance requirements, traceability"
    },
    "communication": {
        "name": "Communication & Signals Standards",
        "keywords": ["hart", "fieldbus", "profibus", "modbus", "4-20ma", "signal", "protocol", "communication"],
        "description": "Industrial communication protocols, signal types"
    },
    "condition_monitoring": {
        "name": "Condition Monitoring Standards",
        "keywords": ["vibration", "monitoring", "predictive", "diagnostic", "condition", "health"],
        "description": "Equipment condition monitoring, predictive maintenance"
    },
    "analytical": {
        "name": "Analytical Instrumentation Standards",
        "keywords": ["analyzer", "ph", "conductivity", "oxygen", "gas", "chromatograph", "spectroscopy"],
        "description": "Process analyzers, gas detection, chemical analysis"
    },
    "accessories": {
        "name": "Accessories & Installation Standards",
        "keywords": ["mounting", "installation", "wiring", "cable", "enclosure", "protection", "ip rating"],
        "description": "Installation requirements, accessories, enclosures"
    }
}

# Mapping from domain type to actual .docx filename
STANDARDS_DOCX_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "chroma_data", "Standards")

STANDARD_FILES = {
    "safety": "instrumentation_safety_standards.docx",
    "pressure": "instrumentation_pressure_standards.docx",
    "temperature": "instrumentation_temperature_standards.docx",
    "flow": "instrumentation_flow_standards.docx",
    "level": "instrumentation_level_standards.docx",
    "control": "instrumentation_control_systems_standards.docx",
    "valves": "instrumentation_valves_actuators_standards.docx",
    "calibration": "instrumentation_calibration_maintenance_standards.docx",
    "communication": "instrumentation_comm_signal_standards.docx",
    "condition_monitoring": "instrumentation_condition_monitoring_standards.docx",
    "analytical": "instrumentation_analytical_standards.docx",
    "accessories": "instrumentation_accessories_calibration_standards.docx"
}

# LLM Model to use for all Deep Agent nodes
DEEP_AGENT_LLM_MODEL = "gemini-2.5-flash-lite"


# ============================================================================
# ITERATIVE GENERATION CONFIGURATION
# ============================================================================

# Minimum number of specifications that Deep Agent RAG must generate
MIN_STANDARDS_SPECS_COUNT = 30

# Maximum iterations to prevent infinite loops
MAX_STANDARDS_ITERATIONS = 5

# Additional domains to try if minimum not reached
FALLBACK_DOMAINS_ORDER = [
    "safety", "calibration", "communication", "accessories",
    "pressure", "temperature", "flow", "level", "control",
    "valves", "condition_monitoring", "analytical"
]

# Parallel processing configuration for Standards Deep Agent
MAX_STANDARDS_PARALLEL_WORKERS = 3  # Number of domains to analyze in parallel


# ============================================================================
# STATE DEFINITIONS
# ============================================================================

class StandardConstraint(TypedDict):
    """Single constraint extracted from a standard document"""
    constraint_type: str  # "requirement", "recommendation", "warning", "specification"
    description: str
    value: Optional[str]
    condition: Optional[str]  # e.g., "If operating above 50°C"
    standard_reference: Optional[str]  # e.g., "IEC 61511"
    confidence: float


class WorkerResult(TypedDict):
    """Output from a single worker analyzing one standard document"""
    standard_type: str
    standard_name: str
    constraints: List[StandardConstraint]
    specifications: Dict[str, Any]
    warnings: List[str]
    processing_time_ms: int
    success: bool
    error: Optional[str]


class ConsolidatedSpecs(TypedDict):
    """Consolidated specifications from all standards"""
    merged_constraints: List[StandardConstraint]
    merged_specifications: Dict[str, Any]
    conflicts: List[Dict[str, Any]]
    warnings: List[str]
    sources: List[str]


class StandardsDeepAgentState(TypedDict, total=False):
    """State for the Standards Deep Agent Workflow"""
    # Input
    user_requirement: str
    session_id: str

    # Inferred specs from database (passed in)
    inferred_specs: Dict[str, Any]

    # Planner output
    relevant_standard_types: List[str]
    planning_reasoning: str

    # Worker outputs (collected from parallel execution)
    parallel_results: Annotated[List[WorkerResult], operator.add]

    # Synthesizer output
    consolidated_specs: Optional[ConsolidatedSpecs]

    # Final merged output
    final_specifications: Dict[str, Any]

    # Metadata
    start_time: float
    processing_time_ms: int
    status: str
    error: Optional[str]

    # Workflow tracking
    current_node: str


# ============================================================================
# PROMPTS
# ============================================================================

PLANNER_PROMPT = """
You are a Standards Planning Agent for industrial instrumentation.

Given a user requirement, determine which standard document domains are relevant.

USER REQUIREMENT:
{user_requirement}

AVAILABLE STANDARD DOMAINS:
{available_domains}

INSTRUCTIONS:
1. Analyze the user requirement carefully
2. Identify which standard domains are relevant to address this requirement
3. Consider safety implications (always include "safety" if there's any hazardous area mention)
4. Be thorough but don't include irrelevant domains

Return ONLY valid JSON:
{{
    "relevant_domains": ["domain1", "domain2", ...],
    "reasoning": "Brief explanation of why these domains are selected"
}}
"""

WORKER_PROMPT = """
You are a Standards Analysis Agent specialized in {standard_type} standards.

Your task is to extract ALL relevant constraints and specifications from the provided standard document that apply to the user's requirement.

USER REQUIREMENT:
{user_requirement}

STANDARD DOCUMENT ({standard_name}):
{document_content}

=== CRITICAL: VALUE FORMAT RULES ===

Return ONLY the technical value - NO descriptions, NO explanations, NO sentences.

CORRECT VALUES: "±0.1%", "IP67", "4-20mA HART", "-40 to +85°C", "SIL 2", "316L SS"
WRONG VALUES: "typically ±0.1%", "IP67 for outdoor use", "depends on application"

FORBIDDEN in values:
- Words: "typically", "usually", "approximately", "may be", "can be", "should be"
- Explanations in parentheses: "(depending on...)", "(for example...)", "(typically...)"
- Sentences or descriptions

=== EXTRACTION INSTRUCTIONS ===

1. Read the document carefully for the user's requirement
2. Extract constraints with CLEAN values only
3. Include specific values, ranges, and tolerances
4. Use null for specifications not found in the document
5. Do NOT make up information not present in the document

Return ONLY valid JSON:
{{
    "constraints": [
        {{
            "constraint_type": "requirement|recommendation|warning|specification",
            "description": "Clear description of the constraint",
            "value": "CLEAN technical value (e.g., 'SIL 2', '4-20mA', '±0.1%')",
            "condition": "Condition under which this applies (if any)",
            "standard_reference": "Referenced standard code (e.g., 'IEC 61511')",
            "confidence": 0.0-1.0
        }}
    ],
    "specifications": {{
        "accuracy": "±0.1%" or null,
        "repeatability": "±0.05%" or null,
        "temperature_range": "-40 to +85°C" or null,
        "ambient_temperature": "-20 to +60°C" or null,
        "humidity_range": "0-95% RH" or null,
        "output_signal": "4-20mA HART" or null,
        "supply_voltage": "24 VDC" or null,
        "power_consumption": "< 5W" or null,
        "protection_rating": "IP67" or null,
        "sil_rating": "SIL 2" or null,
        "hazardous_area_approval": "ATEX Zone 1" or null,
        "material_wetted": "316L SS" or null,
        "material_housing": "Aluminum" or null,
        "process_connection": "1/2 NPT" or null,
        "communication_protocol": "HART, Modbus" or null,
        "response_time": "< 250ms" or null,
        "calibration_interval": "12 months" or null,
        "stability": "±0.1% per year" or null,
        "rangeability": "100:1" or null,
        "weight": "1.5 kg" or null,
        "certifications": "CE, IECEx" or null
    }},
    "warnings": ["List of important warnings or cautions"],
    "summary": "Brief summary of findings"
}}
"""


SYNTHESIZER_PROMPT = """
You are a Standards Synthesis Agent for industrial instrumentation.

Your task is to merge and consolidate constraints from multiple standard documents, identifying any conflicts.

USER REQUIREMENT:
{user_requirement}

WORKER RESULTS:
{worker_results}

=== CRITICAL: VALUE FORMAT RULES ===

All values must be CLEAN technical specifications - NO descriptions or explanations.
CORRECT: "±0.1%", "IP67", "4-20mA HART"
WRONG: "typically ±0.1%", "IP67 suitable for outdoor use"

INSTRUCTIONS:
1. Merge all constraints from different standards
2. Identify and flag any CONFLICTS between standards
3. Prioritize safety-related constraints
4. Remove duplicate constraints
5. Consolidate specifications into a unified set - CLEAN VALUES ONLY
6. Only include specifications that have actual values (exclude nulls)

Return ONLY valid JSON:
{{
    "merged_constraints": [
        {{
            "constraint_type": "requirement|recommendation|warning|specification",
            "description": "Merged/consolidated constraint",
            "value": "CLEAN technical value only",
            "condition": "Condition if any",
            "standard_reference": "Source standard(s)",
            "confidence": 0.0-1.0,
            "source_standards": ["list of contributing standards"]
        }}
    ],
    "merged_specifications": {{
        "accuracy": "±0.1%" or null,
        "repeatability": "±0.05%" or null,
        "temperature_range": "-40 to +85°C" or null,
        "ambient_temperature": "-20 to +60°C" or null,
        "humidity_range": "0-95% RH" or null,
        "output_signal": "4-20mA HART" or null,
        "supply_voltage": "24 VDC" or null,
        "power_consumption": "< 5W" or null,
        "protection_rating": "IP67" or null,
        "sil_rating": "SIL 2" or null,
        "hazardous_area_approval": "ATEX Zone 1" or null,
        "material_wetted": "316L SS" or null,
        "material_housing": "Aluminum" or null,
        "process_connection": "1/2 NPT" or null,
        "communication_protocol": "HART, Modbus" or null,
        "response_time": "< 250ms" or null,
        "calibration_interval": "12 months" or null,
        "stability": "±0.1% per year" or null,
        "rangeability": "100:1" or null,
        "weight": "1.5 kg" or null,
        "certifications": "CE, IECEx" or null
    }},
    "conflicts": [
        {{
            "description": "Description of conflict",
            "standard_a": "First standard",
            "value_a": "Value from first standard",
            "standard_b": "Second standard",
            "value_b": "Value from second standard",
            "resolution": "Suggested resolution or 'REQUIRES_REVIEW'"
        }}
    ],
    "warnings": ["Consolidated warnings"],
    "sources": ["List of all standard types analyzed"]
}}
"""


MERGER_PROMPT = """
You are a Specifications Merger Agent for industrial instrumentation.

Your task is to merge standards-based specifications with database-inferred specifications, with standards taking precedence for safety-critical items.

USER REQUIREMENT:
{user_requirement}

STANDARDS-BASED SPECIFICATIONS:
{standards_specs}

DATABASE-INFERRED SPECIFICATIONS:
{inferred_specs}

INSTRUCTIONS:
1. Start with the database-inferred specifications as baseline
2. Overlay standards-based specifications on top
3. For conflicts, standards specifications override EXCEPT for product-specific details
4. Flag any overrides or conflicts for user review
5. Ensure safety-critical specifications from standards are never overridden
6. Output ONLY specifications that have actual values (not null or empty)

STANDARD SPECIFICATION KEYS TO USE:
- accuracy: Measurement accuracy
- temperature_range: Operating temperature range
- process_temperature: Process medium temperature
- pressure_range: Pressure measurement range
- output_signal: Signal type
- supply_voltage: Power supply voltage
- protection_rating: IP/NEMA rating
- hazardous_area_approval: Zone certification
- sil_rating: Safety integrity level
- material_wetted: Wetted parts material
- material_housing: Housing material
- process_connection: Connection type
- response_time: Sensor response time
- communication_protocol: Communication protocol
- calibration_interval: Recommended calibration frequency
- ambient_temperature: Ambient operating range

Return ONLY valid JSON:
{{
    "final_specifications": {{
        "accuracy": "value if applicable",
        "temperature_range": "value if applicable",
        "output_signal": "value if applicable",
        "protection_rating": "value if applicable",
        "sil_rating": "value if applicable",
        "hazardous_area_approval": "value if applicable",
        "material_wetted": "value if applicable",
        "process_connection": "value if applicable",
        "communication_protocol": "value if applicable",
        "supply_voltage": "value if applicable",
        "response_time": "value if applicable",
        "calibration_interval": "value if applicable"
    }},
    "constraints_applied": [
        {{
            "constraint": "Constraint that was applied",
            "impact": "How it affected the final specs"
        }}
    ],
    "overrides": [
        {{
            "field": "Field name",
            "inferred_value": "Original inferred value",
            "standards_value": "Value from standards",
            "reason": "Why standards value was used"
        }}
    ],
    "warnings": ["Any warnings for the user"],
    "confidence": 0.0-1.0
}}
"""


# ============================================================================
# ITERATIVE WORKER PROMPT - For extracting additional specs from more domains
# ============================================================================

ITERATIVE_WORKER_PROMPT = """
You are a Standards Analysis Agent extracting ADDITIONAL specifications.

Your task is to extract specifications from this standard document that are NOT already in the existing list.

USER REQUIREMENT:
{user_requirement}

STANDARD DOCUMENT ({standard_name}):
{document_content}

=== EXISTING SPECIFICATIONS (DO NOT REPEAT THESE) ===
{existing_specs}

=== CRITICAL: VALUE FORMAT RULES ===

Return ONLY clean technical values - NO descriptions, NO explanations.

CORRECT VALUES: "±0.1%", "IP67", "4-20mA HART", "-40 to +85°C", "SIL 2"
WRONG VALUES: "typically ±0.1%", "IP67 for outdoor use", "depends on application"

=== EXTRACT {specs_needed} NEW SPECIFICATIONS ===

Focus on specifications NOT already covered:
- Physical characteristics: dimensions, weight, mounting options
- Electrical parameters: isolation, surge protection, grounding
- Environmental ratings: vibration, shock, altitude, humidity
- Safety requirements: burst pressure, overpressure, failure modes
- Maintenance: MTBF, service intervals, diagnostic coverage
- Certifications: marine, food-grade, nuclear, railway
- Communication: fieldbus options, wireless, redundancy

Return ONLY valid JSON with NEW specifications:
{{
    "specifications": {{
        "new_spec_key_1": "clean technical value",
        "new_spec_key_2": "clean technical value"
    }},
    "constraints": [
        {{
            "constraint_type": "requirement|recommendation|specification",
            "description": "Clear description",
            "value": "CLEAN technical value",
            "standard_reference": "Referenced standard code",
            "confidence": 0.0-1.0
        }}
    ],
    "summary": "Brief summary of additional findings"
}}

CRITICAL: Do NOT repeat any specifications from the existing list above!
"""


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def load_standard_text(standard_type: str) -> Optional[str]:
    """
    Load the full text content of a standard document directly from .docx file.

    This reads the raw .docx file using python-docx for maximum context fidelity.
    No chunking or vector store queries - direct file access.

    Args:
        standard_type: The type of standard (e.g., "safety", "pressure")

    Returns:
        Full text content of the standard document, or None if not found
    """
    try:
        # Get the filename for this standard type
        filename = STANDARD_FILES.get(standard_type)
        if not filename:
            logger.warning(f"No file mapping for standard type: {standard_type}")
            return None
        
        # Build full path
        filepath = os.path.join(STANDARDS_DOCX_DIR, filename)
        
        if not os.path.exists(filepath):
            logger.warning(f"Standard file not found: {filepath}")
            return None
        
        # Load using python-docx
        logger.info(f"Loading raw .docx file: {filename}")
        doc = DocxDocument(filepath)
        
        # Extract all paragraph text
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n\n".join(paragraphs)
        
        logger.info(f"Loaded {len(paragraphs)} paragraphs/rows from {filename} ({len(full_text)} chars)")
        
        return full_text

    except Exception as e:
        logger.error(f"Error loading standard text for {standard_type}: {e}")
        return None


def format_domains_for_prompt() -> str:
    """Format available domains for the planner prompt."""
    lines = []
    for domain_id, domain_info in STANDARD_DOMAINS.items():
        lines.append(f"- {domain_id}: {domain_info['name']}")
        lines.append(f"  Keywords: {', '.join(domain_info['keywords'][:5])}")
        lines.append(f"  Description: {domain_info['description']}")
        lines.append("")
    return "\n".join(lines)


# ============================================================================
# WORKFLOW NODES
# ============================================================================

def plan_analysis_node(state: StandardsDeepAgentState) -> Dict[str, Any]:
    """
    Planner Node: Determine which standard domains are relevant for the requirement.

    Input: User requirement
    Output: List of relevant standard types to analyze
    """
    logger.info("[Planner] Analyzing requirement to determine relevant standards...")

    try:
        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        prompt = ChatPromptTemplate.from_template(PLANNER_PROMPT)
        parser = JsonOutputParser()
        chain = prompt | llm | parser

        result = chain.invoke({
            "user_requirement": state["user_requirement"],
            "available_domains": format_domains_for_prompt()
        })

        relevant_domains = result.get("relevant_domains", [])
        reasoning = result.get("reasoning", "")

        # Validate domains
        valid_domains = [d for d in relevant_domains if d in STANDARD_DOMAINS]

        if not valid_domains:
            # Fallback: include safety and most likely domain based on keywords
            valid_domains = ["safety"]
            requirement_lower = state["user_requirement"].lower()
            for domain_id, domain_info in STANDARD_DOMAINS.items():
                if any(kw in requirement_lower for kw in domain_info["keywords"]):
                    valid_domains.append(domain_id)
                    if len(valid_domains) >= 3:
                        break

        logger.info(f"[Planner] Selected domains: {valid_domains}")
        logger.info(f"[Planner] Reasoning: {reasoning}")

        return {
            "relevant_standard_types": valid_domains,
            "planning_reasoning": reasoning,
            "current_node": "plan_analysis"
        }

    except Exception as e:
        logger.error(f"[Planner] Error: {e}", exc_info=True)
        # Fallback to safety-focused domains
        return {
            "relevant_standard_types": ["safety", "calibration"],
            "planning_reasoning": f"Error during planning: {e}. Using default domains.",
            "error": str(e),
            "current_node": "plan_analysis"
        }


def route_to_workers(state: StandardsDeepAgentState) -> List[Send]:
    """
    Router: Create parallel Send operations for each relevant standard.

    This implements the "map" phase of Map-Reduce by spawning parallel workers.
    """
    logger.info("[Router] Spawning parallel workers for standard analysis...")

    sends = []
    for standard_type in state.get("relevant_standard_types", []):
        sends.append(
            Send(
                "analyze_standard_doc",
                {
                    "standard_type": standard_type,
                    "user_requirement": state["user_requirement"]
                }
            )
        )

    logger.info(f"[Router] Created {len(sends)} parallel worker tasks")
    return sends


def analyze_standard_doc_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Worker Node: Analyze a single standard document.

    This node is executed in parallel for each relevant standard.

    Input: standard_type, user_requirement
    Output: WorkerResult with extracted constraints
    """
    standard_type = state.get("standard_type", "unknown")
    user_requirement = state.get("user_requirement", "")

    logger.info(f"[Worker:{standard_type}] Starting analysis...")
    start_time = time.time()

    try:
        # Load full document text
        document_content = load_standard_text(standard_type)

        if not document_content:
            logger.warning(f"[Worker:{standard_type}] No document content found")
            return {
                "parallel_results": [{
                    "standard_type": standard_type,
                    "standard_name": STANDARD_DOMAINS.get(standard_type, {}).get("name", standard_type),
                    "constraints": [],
                    "specifications": {},
                    "warnings": [f"No document content available for {standard_type}"],
                    "processing_time_ms": int((time.time() - start_time) * 1000),
                    "success": False,
                    "error": "Document not found"
                }]
            }

        # Analyze with LLM
        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        prompt = ChatPromptTemplate.from_template(WORKER_PROMPT)
        parser = JsonOutputParser()
        chain = prompt | llm | parser

        result = chain.invoke({
            "standard_type": standard_type,
            "standard_name": STANDARD_DOMAINS.get(standard_type, {}).get("name", standard_type),
            "user_requirement": user_requirement,
            "document_content": document_content[:30000]  # Limit to ~30k chars for context
        })

        processing_time = int((time.time() - start_time) * 1000)

        worker_result: WorkerResult = {
            "standard_type": standard_type,
            "standard_name": STANDARD_DOMAINS.get(standard_type, {}).get("name", standard_type),
            "constraints": result.get("constraints", []),
            "specifications": result.get("specifications", {}),
            "warnings": result.get("warnings", []),
            "processing_time_ms": processing_time,
            "success": True,
            "error": None
        }

        logger.info(f"[Worker:{standard_type}] Completed in {processing_time}ms, found {len(worker_result['constraints'])} constraints")

        return {"parallel_results": [worker_result]}

    except Exception as e:
        logger.error(f"[Worker:{standard_type}] Error: {e}", exc_info=True)
        return {
            "parallel_results": [{
                "standard_type": standard_type,
                "standard_name": STANDARD_DOMAINS.get(standard_type, {}).get("name", standard_type),
                "constraints": [],
                "specifications": {},
                "warnings": [f"Error analyzing {standard_type}: {str(e)}"],
                "processing_time_ms": int((time.time() - start_time) * 1000),
                "success": False,
                "error": str(e)
            }]
        }


def synthesize_specs_node(state: StandardsDeepAgentState) -> Dict[str, Any]:
    """
    Synthesizer Node: Merge all worker outputs into consolidated specifications.

    This implements the "reduce" phase of Map-Reduce.

    Input: All parallel_results from workers
    Output: ConsolidatedSpecs with merged constraints and conflict detection
    """
    logger.info("[Synthesizer] Merging worker results...")

    parallel_results = state.get("parallel_results", [])

    if not parallel_results:
        logger.warning("[Synthesizer] No worker results to synthesize")
        return {
            "consolidated_specs": {
                "merged_constraints": [],
                "merged_specifications": {},
                "conflicts": [],
                "warnings": ["No standards analysis results available"],
                "sources": []
            },
            "current_node": "synthesize_specs"
        }

    try:
        # Format worker results for prompt
        worker_results_text = json.dumps(parallel_results, indent=2)

        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        prompt = ChatPromptTemplate.from_template(SYNTHESIZER_PROMPT)
        parser = JsonOutputParser()
        chain = prompt | llm | parser

        result = chain.invoke({
            "user_requirement": state["user_requirement"],
            "worker_results": worker_results_text[:50000]  # Limit context
        })

        consolidated: ConsolidatedSpecs = {
            "merged_constraints": result.get("merged_constraints", []),
            "merged_specifications": result.get("merged_specifications", {}),
            "conflicts": result.get("conflicts", []),
            "warnings": result.get("warnings", []),
            "sources": result.get("sources", [])
        }

        logger.info(f"[Synthesizer] Merged {len(consolidated['merged_constraints'])} constraints, detected {len(consolidated['conflicts'])} conflicts")

        return {
            "consolidated_specs": consolidated,
            "current_node": "synthesize_specs"
        }

    except Exception as e:
        logger.error(f"[Synthesizer] Error: {e}", exc_info=True)

        # Fallback: Simple merge without LLM
        all_constraints = []
        all_specs = {}
        all_warnings = []
        sources = []

        for wr in parallel_results:
            all_constraints.extend(wr.get("constraints", []))
            all_specs.update(wr.get("specifications", {}))
            all_warnings.extend(wr.get("warnings", []))
            sources.append(wr.get("standard_type", "unknown"))

        return {
            "consolidated_specs": {
                "merged_constraints": all_constraints,
                "merged_specifications": all_specs,
                "conflicts": [],
                "warnings": all_warnings + [f"Synthesis error: {e}"],
                "sources": sources
            },
            "current_node": "synthesize_specs"
        }


def merge_with_inferred_node(state: StandardsDeepAgentState) -> Dict[str, Any]:
    """
    Merger Node: Combine standards-based specs with database-inferred specs.

    Standards specifications override inferred specs for safety-critical items.

    Input: consolidated_specs, inferred_specs
    Output: final_specifications
    """
    logger.info("[Merger] Merging with inferred specifications...")

    consolidated = state.get("consolidated_specs", {})
    inferred = state.get("inferred_specs", {})

    if not consolidated and not inferred:
        logger.warning("[Merger] No specifications to merge")
        return {
            "final_specifications": {},
            "current_node": "merge_with_inferred",
            "processing_time_ms": int((time.time() - state.get("start_time", time.time())) * 1000)
        }

    try:
        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        prompt = ChatPromptTemplate.from_template(MERGER_PROMPT)
        parser = JsonOutputParser()
        chain = prompt | llm | parser

        result = chain.invoke({
            "user_requirement": state["user_requirement"],
            "standards_specs": json.dumps(consolidated, indent=2),
            "inferred_specs": json.dumps(inferred, indent=2)
        })

        final_specs = {
            "specifications": result.get("final_specifications", {}),
            "constraints_applied": result.get("constraints_applied", []),
            "overrides": result.get("overrides", []),
            "warnings": result.get("warnings", []),
            "confidence": result.get("confidence", 0.0),
            "sources": {
                "standards_analyzed": consolidated.get("sources", []),
                "conflicts_detected": consolidated.get("conflicts", [])
            }
        }

        processing_time = int((time.time() - state.get("start_time", time.time())) * 1000)

        logger.info(f"[Merger] Final specifications generated with confidence: {final_specs['confidence']}")
        logger.info(f"[Merger] Total processing time: {processing_time}ms")

        return {
            "final_specifications": final_specs,
            "processing_time_ms": processing_time,
            "status": "success",
            "current_node": "merge_with_inferred"
        }

    except Exception as e:
        logger.error(f"[Merger] Error: {e}", exc_info=True)

        # Fallback: Return consolidated specs directly
        processing_time = int((time.time() - state.get("start_time", time.time())) * 1000)

        return {
            "final_specifications": {
                "specifications": consolidated.get("merged_specifications", {}),
                "constraints_applied": consolidated.get("merged_constraints", []),
                "overrides": [],
                "warnings": consolidated.get("warnings", []) + [f"Merge error: {e}"],
                "confidence": 0.5,
                "sources": {
                    "standards_analyzed": consolidated.get("sources", []),
                    "conflicts_detected": consolidated.get("conflicts", [])
                }
            },
            "processing_time_ms": processing_time,
            "status": "partial",
            "error": str(e),
            "current_node": "merge_with_inferred"
        }


# ============================================================================
# WORKFLOW GRAPH CONSTRUCTION
# ============================================================================

def create_standards_deep_agent_workflow():
    """
    Create the Standards Deep Agent workflow graph.

    Graph structure (Map-Reduce):

    START → plan_analysis → [route_to_workers]
                                    ↓
                        ┌───────────┼───────────┐
                        ↓           ↓           ↓
                   [Worker 1] [Worker 2] [Worker N]  (PARALLEL)
                        ↓           ↓           ↓
                        └───────────┼───────────┘
                                    ↓
                            synthesize_specs
                                    ↓
                            merge_with_inferred
                                    ↓
                                   END
    """
    logger.info("Creating Standards Deep Agent workflow...")

    # Create graph
    workflow = StateGraph(StandardsDeepAgentState)

    # Add nodes
    workflow.add_node("plan_analysis", plan_analysis_node)
    workflow.add_node("analyze_standard_doc", analyze_standard_doc_node)
    workflow.add_node("synthesize_specs", synthesize_specs_node)
    workflow.add_node("merge_with_inferred", merge_with_inferred_node)

    # Set entry point
    workflow.set_entry_point("plan_analysis")

    # Add conditional edge for parallel worker spawning (Map phase)
    workflow.add_conditional_edges(
        "plan_analysis",
        route_to_workers,
        ["analyze_standard_doc"]
    )

    # Workers converge to synthesizer (Reduce phase)
    workflow.add_edge("analyze_standard_doc", "synthesize_specs")

    # Synthesizer to Merger
    workflow.add_edge("synthesize_specs", "merge_with_inferred")

    # Merger to END
    workflow.add_edge("merge_with_inferred", END)

    # Compile with checkpointing
    app = compile_with_checkpointing(workflow)

    logger.info("Standards Deep Agent workflow created successfully")

    return app


# ============================================================================
# WORKFLOW EXECUTION
# ============================================================================

# Singleton workflow instance
_standards_deep_agent_app = None


def get_standards_deep_agent_workflow():
    """Get or create the Standards Deep Agent workflow instance (singleton)."""
    global _standards_deep_agent_app
    if _standards_deep_agent_app is None:
        _standards_deep_agent_app = create_standards_deep_agent_workflow()
    return _standards_deep_agent_app


def create_standards_deep_agent_state(
    user_requirement: str,
    session_id: Optional[str] = None,
    inferred_specs: Optional[Dict[str, Any]] = None
) -> StandardsDeepAgentState:
    """Create initial state for Standards Deep Agent workflow."""
    return StandardsDeepAgentState(
        user_requirement=user_requirement,
        session_id=session_id or f"deep-{int(time.time())}",
        inferred_specs=inferred_specs or {},
        relevant_standard_types=[],
        planning_reasoning="",
        parallel_results=[],
        consolidated_specs=None,
        final_specifications={},
        start_time=time.time(),
        processing_time_ms=0,
        status="started",
        error=None,
        current_node="start"
    )


def _count_valid_specs(specs_dict: Dict[str, Any]) -> int:
    """
    Count valid (non-null, non-empty) specifications.

    Args:
        specs_dict: Dictionary of specifications

    Returns:
        Count of valid specifications
    """
    if not specs_dict:
        return 0

    count = 0
    for key, value in specs_dict.items():
        if value is not None:
            str_value = str(value).lower().strip()
            if str_value and str_value not in ["null", "none", "n/a", ""]:
                count += 1

    return count


def _extract_additional_specs_from_domain(
    user_requirement: str,
    domain: str,
    existing_specs: Dict[str, Any],
    specs_needed: int
) -> Dict[str, Any]:
    """
    Extract additional specifications from a single domain.

    Args:
        user_requirement: User's requirement string
        domain: Domain to analyze
        existing_specs: Already extracted specifications
        specs_needed: Number of additional specs needed

    Returns:
        Dict with new specifications and constraints
    """
    logger.info(f"[ITERATIVE] Extracting additional specs from domain: {domain}")

    try:
        # Load document
        document_content = load_standard_text(domain)
        if not document_content:
            logger.warning(f"[ITERATIVE] No document found for domain: {domain}")
            return {"specifications": {}, "constraints": []}

        # Format existing specs
        existing_specs_list = "\n".join([
            f"- {key}: {value}"
            for key, value in existing_specs.items()
            if value and str(value).lower() not in ["null", "none", "n/a"]
        ])

        # Create LLM and prompt
        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        prompt = ChatPromptTemplate.from_template(ITERATIVE_WORKER_PROMPT)
        parser = JsonOutputParser()
        chain = prompt | llm | parser

        result = chain.invoke({
            "user_requirement": user_requirement,
            "standard_name": STANDARD_DOMAINS.get(domain, {}).get("name", domain),
            "document_content": document_content[:25000],
            "existing_specs": existing_specs_list,
            "specs_needed": specs_needed
        })

        new_specs = result.get("specifications", {})
        new_constraints = result.get("constraints", [])

        logger.info(f"[ITERATIVE] Domain {domain}: Found {len(new_specs)} new specs")

        return {
            "specifications": new_specs,
            "constraints": new_constraints,
            "domain": domain
        }

    except Exception as e:
        logger.error(f"[ITERATIVE] Error extracting from domain {domain}: {e}")
        return {"specifications": {}, "constraints": [], "error": str(e)}


def _extract_specs_from_domains_parallel(
    user_requirement: str,
    domains: List[str],
    existing_specs: Dict[str, Any],
    specs_needed: int
) -> Dict[str, Dict[str, Any]]:
    """
    Extract specifications from multiple domains IN PARALLEL.

    PARALLELIZATION STRATEGY:
    - Instead of: 3 domains × 1 sequential call = 3 calls
    - Now: 3 parallel workers = ~1.2 calls (estimated with overhead)
    - SPEEDUP: ~2.5x faster

    Args:
        user_requirement: User requirement string
        domains: List of domains to analyze in parallel
        existing_specs: Already extracted specifications
        specs_needed: Target specs per domain

    Returns:
        Dict mapping domain to extracted specs result
    """
    results = {}

    if not domains:
        return results

    logger.info(f"[PARALLEL-DOMAINS] Extracting from {len(domains)} domains IN PARALLEL...")

    # Execute domain extraction in parallel
    with ThreadPoolExecutor(max_workers=min(MAX_STANDARDS_PARALLEL_WORKERS, len(domains))) as executor:
        future_to_domain = {
            executor.submit(
                _extract_additional_specs_from_domain,
                user_requirement, domain, existing_specs, specs_needed
            ): domain for domain in domains
        }

        for future in as_completed(future_to_domain):
            domain = future_to_domain[future]
            try:
                result = future.result()
                results[domain] = result
                logger.info(f"[PARALLEL-DOMAINS] Domain '{domain}' completed with {len(result.get('specifications', {}))} specs")
            except Exception as exc:
                logger.error(f"[PARALLEL-DOMAINS] Domain '{domain}' generated exception: {exc}")
                results[domain] = {"specifications": {}, "constraints": [], "error": str(exc)}

    logger.info(f"[PARALLEL-DOMAINS] Completed parallel extraction from {len(results)} domains")
    return results


def run_standards_deep_agent(
    user_requirement: str,
    session_id: Optional[str] = None,
    inferred_specs: Optional[Dict[str, Any]] = None,
    min_specs: Optional[int] = None
) -> Dict[str, Any]:
    """
    Run the Standards Deep Agent workflow with ITERATIVE LOOP.

    This is the main entry point for the deep agent. It performs parallel
    analysis of relevant standard documents and merges with database-inferred specs.

    ITERATIVE GENERATION: If the initial run produces fewer than MIN_STANDARDS_SPECS_COUNT (30)
    specifications, it will iterate through additional domains until the minimum is reached.

    Args:
        user_requirement: The user's requirement string
        session_id: Optional session ID for checkpointing
        inferred_specs: Optional dict of specifications inferred from database
        min_specs: Optional minimum specs count (defaults to MIN_STANDARDS_SPECS_COUNT)

    Returns:
        Final result dictionary with specifications and metadata
    """
    min_required = min_specs if min_specs is not None else MIN_STANDARDS_SPECS_COUNT
    logger.info(f"Running Standards Deep Agent for: {user_requirement[:100]}... (minimum: {min_required})")

    start_time = time.time()
    all_specifications: Dict[str, Any] = {}
    all_constraints: List[Dict[str, Any]] = []
    all_domains_analyzed: List[str] = []
    iteration_notes: List[str] = []
    iteration = 0

    # Create initial state
    initial_state = create_standards_deep_agent_state(
        user_requirement=user_requirement,
        session_id=session_id,
        inferred_specs=inferred_specs
    )

    # Get workflow
    app = get_standards_deep_agent_workflow()

    # =================================================================
    # ITERATION 1: Initial workflow execution
    # =================================================================
    try:
        iteration = 1
        logger.info(f"[DEEP_AGENT] Iteration {iteration}: Initial workflow execution...")

        config = {"configurable": {"thread_id": session_id or "default"}}
        final_state = app.invoke(initial_state, config=config)

        logger.info(f"Deep Agent iteration {iteration} completed with status: {final_state.get('status', 'unknown')}")

        # Extract specifications from final state
        final_specs = final_state.get("final_specifications", {})
        if isinstance(final_specs, dict) and "specifications" in final_specs:
            all_specifications.update(final_specs.get("specifications", {}))
        elif isinstance(final_specs, dict):
            all_specifications.update(final_specs)

        # Also get specs from consolidated specs
        consolidated = final_state.get("consolidated_specs", {})
        if isinstance(consolidated, dict) and "merged_specifications" in consolidated:
            merged_specs = consolidated.get("merged_specifications", {})
            for key, value in merged_specs.items():
                if key not in all_specifications and value and str(value).lower() not in ["null", "none"]:
                    all_specifications[key] = value

        # Track constraints
        if consolidated and "merged_constraints" in consolidated:
            all_constraints.extend(consolidated.get("merged_constraints", []))

        # Track domains analyzed
        all_domains_analyzed = final_state.get("relevant_standard_types", [])

        initial_count = _count_valid_specs(all_specifications)
        iteration_notes.append(f"Iteration 1: Generated {initial_count} specs from domains {all_domains_analyzed}")

        logger.info(f"[DEEP_AGENT] Iteration {iteration}: Generated {initial_count} specs, domains: {all_domains_analyzed}")

        # =================================================================
        # ITERATIVE LOOP: Continue until minimum is reached
        # =================================================================
        while _count_valid_specs(all_specifications) < min_required and iteration < MAX_STANDARDS_ITERATIONS:
            iteration += 1
            current_count = _count_valid_specs(all_specifications)
            specs_needed = min(15, min_required - current_count + 5)  # Request a few extra

            logger.info(f"[DEEP_AGENT] Iteration {iteration}: Need {min_required - current_count} more specs...")

            # Find domains not yet analyzed
            remaining_domains = [
                d for d in FALLBACK_DOMAINS_ORDER
                if d not in all_domains_analyzed
            ]

            if not remaining_domains:
                logger.warning(f"[DEEP_AGENT] Iteration {iteration}: No more domains to analyze")
                iteration_notes.append(f"Iteration {iteration}: Stopped - all domains exhausted")
                break

            # Analyze next domain(s) IN PARALLEL
            domains_to_try = remaining_domains[:MAX_STANDARDS_PARALLEL_WORKERS]  # Try up to 3 domains in parallel
            logger.info(f"[DEEP_AGENT] Iteration {iteration}: Using PARALLEL extraction for domains: {domains_to_try}")

            # Extract from multiple domains in parallel for speedup
            parallel_results = _extract_specs_from_domains_parallel(
                user_requirement=user_requirement,
                domains=domains_to_try,
                existing_specs=all_specifications,
                specs_needed=specs_needed
            )

            added_count = 0
            for domain, result in parallel_results.items():
                try:
                    new_specs = result.get("specifications", {})
                    new_constraints = result.get("constraints", [])

                    # Add new specs (avoid duplicates)
                    for key, value in new_specs.items():
                        normalized_key = key.lower().replace(" ", "_").replace("-", "_")
                        existing_keys = {k.lower().replace(" ", "_").replace("-", "_") for k in all_specifications.keys()}

                        if normalized_key not in existing_keys and value and str(value).lower() not in ["null", "none"]:
                            all_specifications[key] = value
                            added_count += 1

                    # Add constraints
                    all_constraints.extend(new_constraints)
                    if domain not in all_domains_analyzed:
                        all_domains_analyzed.append(domain)

                except Exception as domain_error:
                    logger.error(f"[DEEP_AGENT] Iteration {iteration}: Error with domain {domain}: {domain_error}")

            iteration_notes.append(f"Iteration {iteration}: Added {added_count} new specs from domains {domains_to_try}")
            logger.info(f"[DEEP_AGENT] Iteration {iteration}: Added {added_count} new specs, total: {_count_valid_specs(all_specifications)}")

            # If no new specs were added, break to avoid infinite loop
            if added_count == 0:
                logger.warning(f"[DEEP_AGENT] Iteration {iteration}: No new specs added, stopping iterations")
                iteration_notes.append(f"Iteration {iteration}: Stopped - no new specs could be extracted")
                break

        # =================================================================
        # FINAL RESULT
        # =================================================================
        final_count = _count_valid_specs(all_specifications)
        target_reached = final_count >= min_required
        processing_time = int((time.time() - start_time) * 1000)

        if target_reached:
            logger.info(f"[DEEP_AGENT] ✓ Target reached: {final_count} specs (minimum: {min_required}) in {iteration} iterations")
        else:
            logger.warning(f"[DEEP_AGENT] ✗ Target NOT reached: {final_count} specs (minimum: {min_required}) after {iteration} iterations")

        return {
            "success": True,
            "status": "completed" if target_reached else "partial",
            "final_specifications": {
                "specifications": all_specifications,
                "constraints_applied": all_constraints,
                "confidence": 0.85 if target_reached else 0.6
            },
            "consolidated_specs": final_state.get("consolidated_specs", {}),
            "standards_analyzed": all_domains_analyzed,
            "planning_reasoning": final_state.get("planning_reasoning", ""),
            "worker_results": final_state.get("parallel_results", []),
            "processing_time_ms": processing_time,
            "iterations": iteration,
            "specs_count": final_count,
            "min_required": min_required,
            "target_reached": target_reached,
            "iteration_notes": "; ".join(iteration_notes),
            "error": final_state.get("error")
        }

    except Exception as e:
        logger.error(f"Error executing Standards Deep Agent: {e}", exc_info=True)
        processing_time = int((time.time() - start_time) * 1000)
        final_count = _count_valid_specs(all_specifications)

        return {
            "success": False,
            "status": "error",
            "error": str(e),
            "final_specifications": {
                "specifications": all_specifications,
                "constraints_applied": all_constraints,
                "confidence": 0.3
            },
            "consolidated_specs": {},
            "standards_analyzed": all_domains_analyzed,
            "processing_time_ms": processing_time,
            "iterations": iteration,
            "specs_count": final_count,
            "min_required": min_required,
            "target_reached": False,
            "iteration_notes": "; ".join(iteration_notes)
        }


# ============================================================================
# BATCH PROCESSING FOR MULTIPLE ITEMS
# ============================================================================

BATCH_PLANNER_PROMPT = """
You are a Standards Planner Agent for industrial instrumentation.

Your task is to analyze MULTIPLE items and determine which standard domains are relevant for ALL of them collectively.

ITEMS TO ANALYZE:
{items_summary}

AVAILABLE STANDARD DOMAINS:
{available_domains}

INSTRUCTIONS:
1. Identify all unique standard domains needed to cover ALL items
2. Prioritize safety-related domains (safety, hazardous area)
3. Group items by their primary domain requirements
4. Return a consolidated list of domains (max 5 to optimize performance)

Return ONLY valid JSON:
{{
    "relevant_domains": ["domain1", "domain2", ...],
    "reasoning": "Why these domains were selected",
    "item_domain_mapping": {{
        "item_name_1": ["domain1", "domain2"],
        "item_name_2": ["domain2", "domain3"]
    }}
}}
"""

BATCH_WORKER_PROMPT = """
You are a Standards Extraction Worker for industrial instrumentation.

Your task is to analyze a standard document and extract specifications for MULTIPLE items at once.

STANDARD DOMAIN: {standard_type}
STANDARD NAME: {standard_name}

ITEMS REQUIRING ANALYSIS:
{items_requiring_this_standard}

DOCUMENT CONTENT:
{document_content}

INSTRUCTIONS:
1. For EACH item in the list, extract relevant specifications from this standard
2. Focus on safety-critical parameters first
3. Include applicable standards references (IEC, ISO, API, etc.)
4. Skip items that are not relevant to this standard domain

STANDARD SPECIFICATION KEYS TO EXTRACT:
- accuracy: Measurement accuracy
- temperature_range: Operating temperature range
- process_temperature: Process medium temperature
- pressure_range: Pressure measurement range
- output_signal: Signal type (4-20mA, HART, etc.)
- supply_voltage: Power supply voltage
- protection_rating: IP/NEMA rating
- hazardous_area_approval: Zone certification (ATEX, IECEx)
- sil_rating: Safety integrity level rating
- material_wetted: Wetted parts material
- material_housing: Housing material
- process_connection: Connection type
- response_time: Sensor response time
- communication_protocol: Communication protocol
- calibration_interval: Calibration frequency
- ambient_temperature: Ambient operating range

Return ONLY valid JSON:
{{
    "items_results": [
        {{
            "item_name": "Name of the item",
            "applicable": true/false,
            "constraints": [
                {{
                    "constraint_type": "requirement|recommendation|warning|specification",
                    "description": "Clear description",
                    "value": "Specific value",
                    "standard_reference": "Referenced standard code"
                }}
            ],
            "specifications": {{
                "accuracy": "extracted value or null",
                "temperature_range": "extracted value or null",
                "output_signal": "extracted value or null",
                "protection_rating": "extracted value or null",
                "sil_rating": "extracted value or null",
                "hazardous_area_approval": "extracted value or null"
            }}
        }}
    ],
    "standard_summary": "Brief summary of this standard's relevance",
    "warnings": ["Any important warnings"]
}}
"""

BATCH_SYNTHESIZER_PROMPT = """
You are a Batch Synthesizer Agent for industrial instrumentation.

Your task is to merge results from multiple standard analyses into final specifications for EACH item.

ITEMS LIST:
{items_list}

WORKER RESULTS FROM ALL STANDARDS:
{worker_results}

INSTRUCTIONS:
1. For EACH item, consolidate specifications from all applicable standards
2. Resolve conflicts by prioritizing safety standards
3. Ensure each item has a complete specification set
4. Mark confidence based on how many standards contributed

Return ONLY valid JSON:
{{
    "items_final_specs": [
        {{
            "item_name": "Item name",
            "item_index": 0,
            "specifications": {{
                "accuracy": "consolidated value or null",
                "temperature_range": "consolidated value or null",
                "output_signal": "consolidated value or null",
                "protection_rating": "consolidated value or null",
                "sil_rating": "consolidated value or null",
                "hazardous_area_approval": "consolidated value or null"
            }},
            "constraints_applied": [
                {{
                    "constraint": "Applied constraint description",
                    "standard_reference": "Source standard"
                }}
            ],
            "standards_analyzed": ["list of standards that contributed"],
            "confidence": 0.0-1.0
        }}
    ],
    "batch_summary": "Overview of the batch analysis",
    "total_standards_analyzed": 3
}}
"""


def _batch_extract_specs_for_items_parallel(
    items_info: List[Dict[str, Any]],
    domain: str,
    document_content: str,
    max_workers: int = 8
) -> Dict[int, Dict[str, Any]]:
    """
    Extract additional specs for multiple items FROM A SINGLE DOMAIN in parallel.
    
    PARALLELIZATION:
    Instead of: 38 sequential LLM calls (one per item)
    Now: 8 parallel workers processing items concurrently
    
    Args:
        items_info: List of item dicts with 'index', 'name', 'existing_specs', 'needed'
        domain: The standard domain to extract from
        document_content: Pre-loaded document content (cached I/O)
        max_workers: Number of parallel workers (default 8)
    
    Returns:
        Dict mapping item index to extracted specs result
    """
    results = {}
    
    if not items_info:
        return results
    
    logger.info(f"[PARALLEL-ITEMS] Processing {len(items_info)} items for domain '{domain}' with {max_workers} workers...")
    
    def extract_for_single_item(item_info: Dict[str, Any]) -> tuple:
        """Worker function to extract specs for one item."""
        item_idx = item_info.get("index", -1)
        item_name = item_info.get("name", f"Item {item_idx}")
        
        try:
            logger.info(f"[ITERATIVE] Extracting additional specs from domain: {domain}")
            
            result = _extract_additional_specs_from_domain(
                user_requirement=item_name,
                domain=domain,
                existing_specs=item_info.get("existing_specs", {}),
                specs_needed=item_info.get("needed", 30)
            )
            
            new_specs = result.get("specifications", {})
            logger.info(f"[ITERATIVE] Domain {domain}: Found {len(new_specs)} new specs")
            
            return (item_idx, result)
            
        except Exception as e:
            logger.error(f"[PARALLEL-ITEMS] Error extracting specs for '{item_name}' from {domain}: {e}")
            return (item_idx, {"specifications": {}, "constraints": [], "error": str(e)})
    
    # Execute parallel extraction using ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=min(max_workers, len(items_info))) as executor:
        future_to_item = {
            executor.submit(extract_for_single_item, item): item 
            for item in items_info
        }
        
        for future in as_completed(future_to_item):
            item = future_to_item[future]
            try:
                idx, result = future.result()
                results[idx] = result
            except Exception as exc:
                item_idx = item.get("index", -1)
                logger.error(f"[PARALLEL-ITEMS] Exception for item index {item_idx}: {exc}")
                results[item_idx] = {"specifications": {}, "constraints": [], "error": str(exc)}
    
    logger.info(f"[PARALLEL-ITEMS] Completed parallel extraction for {len(results)} items from domain '{domain}'")
    return results

def run_standards_deep_agent_batch(
    items: List[Dict[str, Any]],
    session_id: Optional[str] = None,
    domain_context: Optional[str] = None,
    safety_requirements: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Run Standards Deep Agent in BATCH mode for multiple items.
    
    This is the optimized entry point that processes all items in a single
    batch operation instead of making N separate API calls.
    
    EFFICIENCY GAINS:
    - Instead of: N items × 4 LLM calls = 4N calls
    - Now: 1 batch plan + K domain analyses + 1 synthesis = K+2 calls (typically 4-5)
    
    Args:
        items: List of items to enrich, each with 'name', 'category', 'sample_input', 'specifications'
        session_id: Optional session ID for logging
        domain_context: Optional domain context (e.g., "Chemical", "Oil & Gas")
        safety_requirements: Optional safety context (e.g., {"sil_level": "SIL2", "hazardous_area": True})
    
    Returns:
        Dict with enriched items and batch metadata
    """
    start_time = time.time()
    logger.info(f"[BATCH] Starting batch Deep Agent for {len(items)} items...")
    
    if not items:
        return {
            "success": True,
            "items": [],
            "batch_metadata": {
                "total_items": 0,
                "processing_time_ms": 0,
                "standards_analyzed": []
            }
        }
    
    try:
        # ============================================
        # STEP 1: BATCH PLANNING - Single LLM call
        # ============================================
        logger.info("[BATCH] Step 1: Planning - determining relevant domains for all items...")
        
        # Build items summary for planner
        items_summary = []
        for i, item in enumerate(items):
            item_desc = f"{i+1}. {item.get('name', 'Unknown')} ({item.get('category', 'Unknown')})"
            if item.get('sample_input'):
                item_desc += f" - {item.get('sample_input', '')[:100]}"
            items_summary.append(item_desc)
        
        # Add context if available
        context_string = ""
        if domain_context:
            context_string += f"\nDomain Context: {domain_context}"
        if safety_requirements:
            if safety_requirements.get('sil_level'):
                context_string += f"\nSafety Level: {safety_requirements['sil_level']}"
            if safety_requirements.get('hazardous_area'):
                context_string += f"\nHazardous Area: Yes"
        
        items_summary_text = "\n".join(items_summary) + context_string
        
        llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
        
        # Batch planning
        planner_prompt = ChatPromptTemplate.from_template(BATCH_PLANNER_PROMPT)
        planner_parser = JsonOutputParser()
        planner_chain = planner_prompt | llm | planner_parser
        
        plan_result = planner_chain.invoke({
            "items_summary": items_summary_text,
            "available_domains": format_domains_for_prompt()
        })
        
        relevant_domains = plan_result.get("relevant_domains", [])
        item_domain_mapping = plan_result.get("item_domain_mapping", {})
        
        # Validate and limit domains
        valid_domains = [d for d in relevant_domains if d in STANDARD_DOMAINS]
        if not valid_domains:
            valid_domains = ["safety", "accessories"]  # Fallback
        valid_domains = valid_domains[:5]  # Limit to 5 domains max
        
        logger.info(f"[BATCH] Selected {len(valid_domains)} domains: {valid_domains}")
        
        # ============================================
        # STEP 2: PARALLEL DOCUMENT ANALYSIS
        # ============================================
        logger.info(f"[BATCH] Step 2: Analyzing {len(valid_domains)} standard documents IN PARALLEL...")
        
        # Pre-load all documents (cached I/O)
        document_cache = {}
        for domain in valid_domains:
            doc_content = load_standard_text(domain)
            if doc_content:
                document_cache[domain] = doc_content
        
        # Build items list for each domain
        all_worker_results = []
        
        def analyze_single_domain(domain: str) -> Dict[str, Any]:
            """Analyze a single domain - designed to run in parallel."""
            if domain not in document_cache:
                logger.warning(f"[BATCH] Skipping domain {domain} - no document found")
                return None
            
            # Determine which items need this domain
            relevant_items = []
            for i, item in enumerate(items):
                item_name = item.get('name', f'Item {i+1}')
                # Check if this domain is mapped to this item
                if item_domain_mapping.get(item_name):
                    if domain in item_domain_mapping[item_name]:
                        relevant_items.append({
                            "index": i,
                            "name": item_name,
                            "category": item.get('category', ''),
                            "sample_input": item.get('sample_input', '')[:200]
                        })
                else:
                    # Fallback: include all items for safety domain
                    if domain == "safety":
                        relevant_items.append({
                            "index": i,
                            "name": item_name,
                            "category": item.get('category', ''),
                            "sample_input": item.get('sample_input', '')[:200]
                        })
            
            if not relevant_items:
                # Include all items if no specific mapping
                relevant_items = [
                    {
                        "index": i,
                        "name": item.get('name', f'Item {i+1}'),
                        "category": item.get('category', ''),
                        "sample_input": item.get('sample_input', '')[:200]
                    }
                    for i, item in enumerate(items)
                ]
            
            logger.info(f"[BATCH-PARALLEL] Analyzing domain '{domain}' for {len(relevant_items)} items...")
            
            # Create a separate LLM instance for this thread to enable true parallelism
            thread_llm = create_llm_with_fallback(model=DEEP_AGENT_LLM_MODEL, temperature=0.1)
            
            # Single LLM call for this domain covering all relevant items
            worker_prompt = ChatPromptTemplate.from_template(BATCH_WORKER_PROMPT)
            worker_parser = JsonOutputParser()
            worker_chain = worker_prompt | thread_llm | worker_parser
            
            try:
                worker_result = worker_chain.invoke({
                    "standard_type": domain,
                    "standard_name": STANDARD_DOMAINS.get(domain, {}).get("name", domain),
                    "items_requiring_this_standard": json.dumps(relevant_items, indent=2),
                    "document_content": document_cache[domain][:25000]  # Limit context
                })
                
                logger.info(f"[BATCH-PARALLEL] Domain '{domain}' completed - {len(worker_result.get('items_results', []))} item results")
                
                return {
                    "domain": domain,
                    "domain_name": STANDARD_DOMAINS.get(domain, {}).get("name", domain),
                    "items_results": worker_result.get("items_results", []),
                    "warnings": worker_result.get("warnings", [])
                }
                
            except Exception as e:
                logger.error(f"[BATCH-PARALLEL] Error analyzing domain {domain}: {e}")
                return {
                    "domain": domain,
                    "domain_name": STANDARD_DOMAINS.get(domain, {}).get("name", domain),
                    "items_results": [],
                    "warnings": [f"Error: {str(e)}"],
                    "error": str(e)
                }
        
        # Execute domain analysis in parallel using ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=min(len(valid_domains), 5)) as executor:
            future_to_domain = {executor.submit(analyze_single_domain, domain): domain for domain in valid_domains}
            for future in as_completed(future_to_domain):
                domain = future_to_domain[future]
                try:
                    result = future.result()
                    if result:
                        all_worker_results.append(result)
                except Exception as exc:
                    logger.error(f"[BATCH-PARALLEL] Domain {domain} generated an exception: {exc}")
                    all_worker_results.append({
                        "domain": domain,
                        "domain_name": STANDARD_DOMAINS.get(domain, {}).get("name", domain),
                        "items_results": [],
                        "warnings": [f"Exception: {str(exc)}"],
                        "error": str(exc)
                    })
        
        # ============================================
        # STEP 3: BATCH SYNTHESIS - Single LLM call
        # ============================================
        logger.info("[BATCH] Step 3: Synthesizing results for all items...")
        
        items_list = [
            {"index": i, "name": item.get('name', f'Item {i+1}'), "category": item.get('category', '')}
            for i, item in enumerate(items)
        ]
        
        synth_prompt = ChatPromptTemplate.from_template(BATCH_SYNTHESIZER_PROMPT)
        synth_parser = JsonOutputParser()
        synth_chain = synth_prompt | llm | synth_parser
        
        synth_result = synth_chain.invoke({
            "items_list": json.dumps(items_list, indent=2),
            "worker_results": json.dumps(all_worker_results, indent=2)[:50000]  # Limit context
        })
        
        items_final_specs = synth_result.get("items_final_specs", [])
        
        # ============================================
        # STEP 4: MERGE RESULTS BACK INTO ITEMS
        # ============================================
        logger.info("[BATCH] Step 4: Merging specifications into items...")
        
        enriched_items = []
        for i, item in enumerate(items):
            # Find the synthesized specs for this item
            item_specs = None
            for spec in items_final_specs:
                if spec.get("item_index") == i or spec.get("item_name") == item.get("name"):
                    item_specs = spec
                    break
            
            # Create enriched item copy
            enriched_item = dict(item)
            
            if item_specs:
                standards_specs = item_specs.get("specifications", {})
                standards_analyzed = item_specs.get("standards_analyzed", [])
                confidence = item_specs.get("confidence", 0.8)
                
                # Build combined specifications (80% standards, 20% database)
                combined_specs = {}
                
                # Add standards specs (priority)
                for key, value in standards_specs.items():
                    if value and str(value).lower() not in ["null", "none", "n/a", ""]:
                        combined_specs[key] = {
                            "value": value,
                            "source": "standards",
                            "confidence": confidence
                        }
                
                # Add database specs (non-conflicting)
                db_specs = item.get("specifications", {})
                for key, value in db_specs.items():
                    if key not in combined_specs and value:
                        combined_specs[key] = {
                            "value": value,
                            "source": "database",
                            "confidence": 0.7
                        }
                
                # Calculate source breakdown
                standards_count = sum(1 for v in combined_specs.values() if v.get("source") == "standards")
                db_count = sum(1 for v in combined_specs.values() if v.get("source") == "database")
                total_count = len(combined_specs)
                
                standards_pct = round((standards_count / total_count) * 100) if total_count > 0 else 0
                db_pct = round((db_count / total_count) * 100) if total_count > 0 else 0
                
                # Update enriched item
                enriched_item["standards_info"] = {
                    "enrichment_status": "success",
                    "enrichment_method": "batch_deep_agent",
                    "standards_analyzed": standards_analyzed,
                    "constraints_applied": item_specs.get("constraints_applied", []),
                    "confidence": confidence
                }
                enriched_item["standards_specifications"] = standards_specs
                enriched_item["combined_specifications"] = combined_specs
                enriched_item["specification_source"] = {
                    "standards_pct": standards_pct,
                    "database_pct": db_pct,
                    "standards_count": standards_count,
                    "database_count": db_count,
                    "total_count": total_count
                }
                
                logger.info(f"[BATCH] Item '{item.get('name', i)}': {standards_pct}% standards, {db_pct}% database")
            else:
                # No specs found - use database only
                enriched_item["standards_info"] = {
                    "enrichment_status": "no_match",
                    "enrichment_method": "batch_deep_agent",
                    "error": "No standards specifications found for this item"
                }
                enriched_item["specification_source"] = {
                    "standards_pct": 0,
                    "database_pct": 100,
                    "fallback": True
                }
            
            enriched_items.append(enriched_item)

        # ============================================
        # STEP 5: ITERATIVE ENRICHMENT FOR ITEMS BELOW MINIMUM
        # ============================================
        logger.info(f"[BATCH] Step 5: Checking items for minimum spec count ({MIN_STANDARDS_SPECS_COUNT})...")

        # Track which items need more specs
        items_needing_more = []
        for i, enriched_item in enumerate(enriched_items):
            combined_specs = enriched_item.get("combined_specifications", {})
            standards_specs = enriched_item.get("standards_specifications", {})
            all_item_specs = {**combined_specs, **standards_specs}

            spec_count = _count_valid_specs(all_item_specs)
            if spec_count < MIN_STANDARDS_SPECS_COUNT:
                items_needing_more.append({
                    "index": i,
                    "name": enriched_item.get("name", f"Item {i+1}"),
                    "current_count": spec_count,
                    "needed": MIN_STANDARDS_SPECS_COUNT - spec_count,
                    "existing_specs": all_item_specs
                })

        if items_needing_more:
            logger.info(f"[BATCH] {len(items_needing_more)} items need additional specifications")

            # Find domains not yet analyzed
            remaining_domains = [d for d in FALLBACK_DOMAINS_ORDER if d not in valid_domains]
            iteration = 0

            while items_needing_more and remaining_domains and iteration < MAX_STANDARDS_ITERATIONS:
                iteration += 1
                domains_to_try = remaining_domains[:2]  # Try 2 domains per iteration
                remaining_domains = remaining_domains[2:]

                logger.info(f"[BATCH] Iteration {iteration}: Trying additional domains {domains_to_try}")

                for domain in domains_to_try:
                    if domain not in document_cache:
                        doc_content = load_standard_text(domain)
                        if doc_content:
                            document_cache[domain] = doc_content
                        else:
                            continue

                    # ================================================================
                    # PARALLEL PROCESSING: Process ALL items for this domain at once
                    # ================================================================
                    # Instead of sequential: for item_info in items_needing_more[:] (one LLM call per item)
                    # Now parallel: 8 workers processing items concurrently
                    
                    parallel_results = _batch_extract_specs_for_items_parallel(
                        items_info=items_needing_more,
                        domain=domain,
                        document_content=document_cache[domain],
                        max_workers=8  # Process up to 8 items simultaneously
                    )
                    
                    # Merge parallel results back into enriched_items
                    for item_info in items_needing_more[:]:  # Copy list for safe removal
                        item_result = parallel_results.get(item_info["index"], {})
                        new_specs = item_result.get("specifications", {})
                        added_count = 0

                        # Add new specs to the enriched item
                        enriched_item = enriched_items[item_info["index"]]
                        combined_specs = enriched_item.get("combined_specifications", {})
                        standards_specs = enriched_item.get("standards_specifications", {})

                        for key, value in new_specs.items():
                            normalized_key = key.lower().replace(" ", "_").replace("-", "_")
                            existing_keys = {
                                k.lower().replace(" ", "_").replace("-", "_")
                                for k in {**combined_specs, **standards_specs, **item_info["existing_specs"]}.keys()
                            }

                            if normalized_key not in existing_keys and value and str(value).lower() not in ["null", "none"]:
                                standards_specs[key] = value
                                combined_specs[key] = {"value": value, "source": "standards_iterative", "confidence": 0.7}
                                item_info["existing_specs"][key] = value
                                added_count += 1

                        enriched_item["standards_specifications"] = standards_specs
                        enriched_item["combined_specifications"] = combined_specs
                        item_info["current_count"] += added_count
                        item_info["needed"] = MIN_STANDARDS_SPECS_COUNT - item_info["current_count"]

                        # Update standards analyzed
                        if "standards_info" in enriched_item:
                            analyzed = enriched_item["standards_info"].get("standards_analyzed", [])
                            if domain not in analyzed:
                                analyzed.append(domain)
                                enriched_item["standards_info"]["standards_analyzed"] = analyzed

                        logger.info(f"[BATCH] Item '{item_info['name']}': Added {added_count} specs from {domain}, total: {item_info['current_count']}")

                        # Check if item has reached minimum
                        if item_info["current_count"] >= MIN_STANDARDS_SPECS_COUNT:
                            items_needing_more.remove(item_info)
                            logger.info(f"[BATCH] ✓ Item '{item_info['name']}' reached minimum ({item_info['current_count']} specs)")

                # Update the valid_domains for metadata
                valid_domains.extend(domains_to_try)

        # Update spec counts in enriched items
        for enriched_item in enriched_items:
            combined_specs = enriched_item.get("combined_specifications", {})
            standards_specs = enriched_item.get("standards_specifications", {})

            standards_count = sum(1 for v in combined_specs.values() if isinstance(v, dict) and v.get("source", "").startswith("standards"))
            db_count = sum(1 for v in combined_specs.values() if isinstance(v, dict) and v.get("source") == "database")
            total_count = len(combined_specs)

            standards_pct = round((standards_count / total_count) * 100) if total_count > 0 else 0
            db_pct = round((db_count / total_count) * 100) if total_count > 0 else 0

            enriched_item["specification_source"] = {
                "standards_pct": standards_pct,
                "database_pct": db_pct,
                "standards_count": standards_count,
                "database_count": db_count,
                "total_count": total_count,
                "min_required": MIN_STANDARDS_SPECS_COUNT,
                "target_reached": total_count >= MIN_STANDARDS_SPECS_COUNT
            }

        # Calculate totals
        processing_time = int((time.time() - start_time) * 1000)
        successful_count = sum(
            1 for item in enriched_items
            if item.get("standards_info", {}).get("enrichment_status") == "success"
        )
        items_at_minimum = sum(
            1 for item in enriched_items
            if item.get("specification_source", {}).get("target_reached", False)
        )

        logger.info(f"[BATCH] Completed: {successful_count}/{len(items)} items enriched in {processing_time}ms")
        logger.info(f"[BATCH] {items_at_minimum}/{len(items)} items reached minimum spec count ({MIN_STANDARDS_SPECS_COUNT})")
        logger.info(f"[BATCH] Efficiency: {len(valid_domains) + 2} LLM calls instead of {len(items) * 4}")
        
        return {
            "success": True,
            "items": enriched_items,
            "batch_metadata": {
                "total_items": len(items),
                "successful_enrichments": successful_count,
                "items_at_minimum": items_at_minimum,
                "min_specs_required": MIN_STANDARDS_SPECS_COUNT,
                "processing_time_ms": processing_time,
                "standards_analyzed": valid_domains,
                "llm_calls_made": len(valid_domains) + 2,  # 1 planner + N workers + 1 synthesizer
                "llm_calls_saved": (len(items) * 4) - (len(valid_domains) + 2)
            }
        }
        
    except Exception as e:
        logger.error(f"[BATCH] Batch processing failed: {e}", exc_info=True)
        processing_time = int((time.time() - start_time) * 1000)
        
        # Return items unchanged with error info
        for item in items:
            item["standards_info"] = {
                "enrichment_status": "error",
                "enrichment_method": "batch_deep_agent",
                "error": str(e)
            }
        
        return {
            "success": False,
            "items": items,
            "error": str(e),
            "batch_metadata": {
                "total_items": len(items),
                "successful_enrichments": 0,
                "processing_time_ms": processing_time,
                "standards_analyzed": []
            }
        }


# ============================================================================
# TESTING
# ============================================================================

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    logger.info("=" * 80)
    logger.info("TESTING STANDARDS DEEP AGENT")
    logger.info("=" * 80)

    test_requirements = [
        "I need a pressure transmitter for a SIL3 application in Zone 1",
        "What are the specs for a flow meter for hydrogen in a hazardous area?",
        "Temperature sensor for high-temperature furnace with 4-20mA output"
    ]

    # Sample inferred specs from database
    sample_inferred_specs = {
        "product_type": "pressure_transmitter",
        "measurement_range": "0-100 bar",
        "output_signal": "4-20mA HART",
        "accuracy": "0.075%"
    }

    for i, requirement in enumerate(test_requirements, 1):
        logger.info(f"\n[Test {i}] Requirement: {requirement}")
        logger.info("-" * 80)

        result = run_standards_deep_agent(
            user_requirement=requirement,
            session_id=f"test-{i}",
            inferred_specs=sample_inferred_specs if i == 1 else None
        )

        if result["success"]:
            logger.info(f"\nStatus: {result['status']}")
            logger.info(f"Standards Analyzed: {result['standards_analyzed']}")
            logger.info(f"Processing Time: {result['processing_time_ms']}ms")
            logger.info(f"\nPlanning Reasoning: {result['planning_reasoning']}")

            if result.get("final_specifications"):
                logger.info(f"\nFinal Specifications:")
                specs = result["final_specifications"]
                if specs.get("specifications"):
                    for key, value in list(specs["specifications"].items())[:5]:
                        logger.info(f"  {key}: {value}")

                if specs.get("warnings"):
                    logger.info(f"\nWarnings:")
                    for warning in specs["warnings"][:3]:
                        logger.info(f"  - {warning}")
        else:
            logger.error(f"\n[ERROR] {result.get('error', 'Unknown error')}")

        logger.info("=" * 80)
